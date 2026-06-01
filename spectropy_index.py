"""
spectropy_index.py
------------------
Generate a SPECTROPY "TABLE OF CONTENTS" index page as a .docx file,
matching the locked sample format:

  - A4-ish page: 21 cm x 27.8 cm, margins top=0 bottom=0 left=2.57cm right=1.83cm
  - Navy (#002060) left vertical bar
  - Centered navy bold title "TABLE OF CONTENTS" (26pt) + navy underline rule
  - 3-column table: S.NO | Chapter Name | Page. No
  - Navy header row, white bold text (20pt)
  - Chapter rows: auto-numbered S.NO (1, 2, ...), bold chapter name (16pt)
  - Topic rows: bold "1.1" prefix + regular topic text (14pt),
    bold zero-padded page number (18pt), numbering resets per chapter
  - Font: Calibri everywhere (explicit, no theme dependency)

Dependency: python-docx  ->  pip install python-docx

Usage:
    from spectropy_index import generate_index

    data = [
        # (chapter_name, topic_text, page_number)
        ("CELL THE UNIT OF LIFE", "Introduction to cell structure and cell theory", 1),
        ("CELL THE UNIT OF LIFE", "Comparing plant and animal cells", 21),
        ("CELL CYCLE: DIVISION", "Phases of the cell cycle and quiescent stage", 42),
        ("CELL CYCLE: DIVISION", "Mitosis: prophase and metaphase events", 58),
    ]
    generate_index(data, "FINAL_INDEX.docx", title="TABLE OF CONTENTS")
"""

from docx import Document
from docx.shared import Pt, RGBColor, Twips, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ---- Locked style constants (from the SPECTROPY sample, source of truth) ----
NAVY = RGBColor(0x00, 0x20, 0x60)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
FONT = "Calibri"

# Column widths in DXA/twips (sum = 9270)
COL_SNO = 988
COL_NAME = 6662
COL_PAGE = 1620

# Page geometry (DXA/twips). 1 cm = 566.929 twips
PAGE_W = 11906      # 21 cm
PAGE_H = 15761      # 27.8 cm
MARGIN_TOP = 567    # 1 cm
MARGIN_BOTTOM = 567  # 1 cm
MARGIN_LEFT = 1457  # 2.57 cm
MARGIN_RIGHT = 1038  # 1.83 cm


def _force_calibri_body_theme(doc):
    """The stock python-docx theme uses Cambria as the minor (body) font, so a
    theme-linked run would render as "Cambria (Body)". Repoint the theme's
    minor font to Calibri so theme-linked runs show as "Calibri (Body)".
    """
    import re
    part = doc.part
    theme_part = None
    for rel in part.rels.values():
        if rel.reltype.endswith("/theme"):
            theme_part = rel.target_part
            break
    if theme_part is None:
        return
    xml = theme_part.blob.decode("utf-8")
    xml = re.sub(
        r'(<a:minorFont>\s*<a:latin typeface=")[^"]*(")',
        r"\1Calibri\2", xml, count=1,
    )
    theme_part._blob = xml.encode("utf-8")


def _apply_body_font(rfonts):
    """Point an existing <w:rFonts> at the document theme's minor (body) font
    so Word displays it as "Calibri (Body)" rather than a plain "Calibri".
    Any pre-existing explicit names are cleared so the theme wins.
    """
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        if rfonts.get(qn(attr)) is not None:
            del rfonts.attrib[qn(attr)]
    rfonts.set(qn("w:asciiTheme"), "minorHAnsi")
    rfonts.set(qn("w:hAnsiTheme"), "minorHAnsi")
    rfonts.set(qn("w:eastAsiaTheme"), "minorHAnsi")
    rfonts.set(qn("w:cstheme"), "minorBidi")


def _set_cell_bg(cell, hex_color):
    """Apply solid background shading to a table cell."""
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    cell._tc.get_or_add_tcPr().append(shd)


def _set_cell_vertical_center(cell):
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def _set_cell_width(cell, dxa):
    cell.width = Twips(dxa)
    tcW = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
    if tcW is None:
        tcW = OxmlElement("w:tcW")
        cell._tc.get_or_add_tcPr().append(tcW)
    tcW.set(qn("w:w"), str(dxa))
    tcW.set(qn("w:type"), "dxa")


def _run(paragraph, text, size_pt, bold=False, color=None):
    r = paragraph.add_run(text)
    r.font.size = Pt(size_pt)
    r.font.bold = bold
    if color is not None:
        r.font.color.rgb = color
    # use the theme body font everywhere -> shows as "Calibri (Body)"
    rpr = r._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    _apply_body_font(rfonts)
    return r


def _clear_para_spacing(paragraph, before=120, after=120):
    pf = paragraph.paragraph_format
    pf.space_before = Twips(before)
    pf.space_after = Twips(after)


def _add_left_page_bar(doc, section):
    """Add the navy left vertical bar as a filled rectangle SHAPE anchored
    in the header (so it spans the full page height), exactly like the
    SPECTROPY sample. This is a drawing/shape, not a page border.
    """
    # Ensure the section has a header part we can write into.
    header = section.header
    header.is_linked_to_previous = False
    # Use the first paragraph of the header to host the anchored drawing.
    if header.paragraphs:
        p = header.paragraphs[0]
    else:
        p = header.add_paragraph()
    run = p.add_run()

    bar_xml = (
        '<w:drawing xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
        'xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing" '
        'xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main" '
        'xmlns:wps="http://schemas.microsoft.com/office/word/2010/wordprocessingShape" '
        'xmlns:wp14="http://schemas.microsoft.com/office/word/2010/wordprocessingDrawing">'
        '<wp:anchor distT="0" distB="0" distL="0" distR="0" simplePos="0" '
        'relativeHeight="251658240" behindDoc="0" locked="0" layoutInCell="1" allowOverlap="1">'
        '<wp:simplePos x="0" y="0"/>'
        '<wp:positionH relativeFrom="page"><wp:posOffset>-3203</wp:posOffset></wp:positionH>'
        '<wp:positionV relativeFrom="page"><wp:posOffset>-3976</wp:posOffset></wp:positionV>'
        '<wp:extent cx="355600" cy="11099800"/>'
        '<wp:effectExtent l="0" t="0" r="25400" b="25400"/>'
        '<wp:wrapNone/>'
        '<wp:docPr id="1888788892" name="LeftBar"/>'
        '<wp:cNvGraphicFramePr><a:graphicFrameLocks/></wp:cNvGraphicFramePr>'
        '<a:graphic>'
        '<a:graphicData uri="http://schemas.microsoft.com/office/word/2010/wordprocessingShape">'
        '<wps:wsp><wps:cNvSpPr><a:spLocks/></wps:cNvSpPr>'
        '<wps:spPr><a:xfrm><a:off x="0" y="0"/><a:ext cx="355600" cy="11099800"/></a:xfrm>'
        '<a:custGeom><a:avLst/><a:gdLst/><a:ahLst/><a:cxnLst/>'
        '<a:rect l="l" t="t" r="r" b="b"/>'
        '<a:pathLst><a:path w="502920" h="10058400">'
        '<a:moveTo><a:pt x="502920" y="10058400"/></a:moveTo>'
        '<a:lnTo><a:pt x="0" y="10058400"/></a:lnTo>'
        '<a:lnTo><a:pt x="0" y="0"/></a:lnTo>'
        '<a:lnTo><a:pt x="502920" y="0"/></a:lnTo>'
        '<a:lnTo><a:pt x="502920" y="10058400"/></a:lnTo>'
        '<a:close/></a:path></a:pathLst></a:custGeom>'
        '<a:solidFill><a:srgbClr val="002060"/></a:solidFill>'
        '<a:ln><a:solidFill><a:srgbClr val="002060"/></a:solidFill></a:ln>'
        '</wps:spPr>'
        '<wps:bodyPr wrap="square" lIns="0" tIns="0" rIns="0" bIns="0" rtlCol="0">'
        '<a:prstTxWarp prst="textNoShape"><a:avLst/></a:prstTxWarp><a:noAutofit/>'
        '</wps:bodyPr></wps:wsp>'
        '</a:graphicData></a:graphic>'
        '<wp14:sizeRelH relativeFrom="margin"><wp14:pctWidth>0</wp14:pctWidth></wp14:sizeRelH>'
        '<wp14:sizeRelV relativeFrom="margin"><wp14:pctHeight>0</wp14:pctHeight></wp14:sizeRelV>'
        '</wp:anchor></w:drawing>'
    )
    from docx.oxml import parse_xml
    run._r.append(parse_xml(bar_xml))


def _add_title_rule(doc):
    """Centered navy bold title + navy underline rule beneath it."""
    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    _run(p, "TABLE OF CONTENTS", 26, bold=True, color=NAVY)

    # Underline rule (paragraph bottom border, navy)
    rule = doc.add_paragraph()
    rule.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rule.paragraph_format.space_after = Pt(14)
    pPr = rule._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "18")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "002060")
    pbdr.append(bottom)
    pPr.append(pbdr)


def _set_table_grid(table, widths):
    """Force fixed column widths via the w:tblGrid element."""
    tbl = table._tbl
    # remove existing grid
    existing = tbl.find(qn("w:tblGrid"))
    if existing is not None:
        tbl.remove(existing)
    grid = OxmlElement("w:tblGrid")
    for w in widths:
        gc = OxmlElement("w:gridCol")
        gc.set(qn("w:w"), str(w))
        grid.append(gc)
    # tblGrid must come right after tblPr
    tblPr = tbl.find(qn("w:tblPr"))
    tblPr.addnext(grid)


def _set_table_width(table, total):
    tblPr = table._tbl.tblPr
    tblW = tblPr.find(qn("w:tblW"))
    if tblW is None:
        tblW = OxmlElement("w:tblW")
        tblPr.append(tblW)
    tblW.set(qn("w:w"), str(total))
    tblW.set(qn("w:type"), "dxa")


def generate_index(data, output_path, title="TABLE OF CONTENTS",
                   zero_pad=True, custom_title_text=None):
    """
    Build the index DOCX.

    Parameters
    ----------
    data : list of (chapter_name, topic_text, page_number) tuples, in order.
           Consecutive rows with the same chapter_name are grouped under one
           chapter row. page_number may be int or str.
    output_path : path to write the .docx file.
    title : kept for backwards compatibility (ignored if custom_title_text set).
    zero_pad : if True, page numbers shown as 01, 09, 15 (min 2 digits).
    custom_title_text : override the title text shown at top.

    Returns
    -------
    output_path
    """
    doc = Document()
    _force_calibri_body_theme(doc)

    # ---- document default font = Calibri (Body) (theme minor font) ----
    normal = doc.styles["Normal"]
    normal.font.size = Pt(11)
    rpr = normal.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    _apply_body_font(rfonts)

    # ---- page geometry ----
    sec = doc.sections[0]
    sec.page_width = Twips(PAGE_W)
    sec.page_height = Twips(PAGE_H)
    sec.top_margin = Twips(MARGIN_TOP)
    sec.bottom_margin = Twips(MARGIN_BOTTOM)
    sec.left_margin = Twips(MARGIN_LEFT)
    sec.right_margin = Twips(MARGIN_RIGHT)
    _add_left_page_bar(doc, sec)

    # ---- title + rule ----
    _add_title_rule(doc)
    if custom_title_text:
        # replace the title run text
        doc.paragraphs[0].runs[0].text = custom_title_text

    # ---- table ----
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = False
    # fixed layout
    tblPr = table._tbl.tblPr
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tblPr.append(layout)

    widths = [COL_SNO, COL_NAME, COL_PAGE]
    _set_table_width(table, sum(widths))
    _set_table_grid(table, widths)

    # ---- header row ----
    hdr = table.rows[0].cells
    hdr_labels = ["S.NO", "Chapter Name", "Page. No"]
    for i, cell in enumerate(hdr):
        _set_cell_width(cell, widths[i])
        _set_cell_bg(cell, "002060")
        _set_cell_vertical_center(cell)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _clear_para_spacing(p)
        _run(p, hdr_labels[i], 18, bold=True, color=WHITE)
    # taller header row
    table.rows[0].height = Twips(736)

    # ---- body rows ----
    last_chapter = None
    chap_no = 0
    topic_no = 0

    for chapter, topic, page in data:
        if chapter != last_chapter:
            chap_no += 1
            topic_no = 0
            last_chapter = chapter
            # chapter row
            row = table.add_row().cells
            _set_cell_width(row[0], widths[0]); _set_cell_vertical_center(row[0])
            _set_cell_width(row[1], widths[1]); _set_cell_vertical_center(row[1])
            _set_cell_width(row[2], widths[2]); _set_cell_vertical_center(row[2])

            p0 = row[0].paragraphs[0]; p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _clear_para_spacing(p0)
            _run(p0, f"{chap_no}.", 16, bold=True)

            p1 = row[1].paragraphs[0]; _clear_para_spacing(p1)
            _run(p1, chapter, 16, bold=True)

            p2 = row[2].paragraphs[0]; p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _clear_para_spacing(p2)  # page cell left blank for chapter rows

        # topic row
        topic_no += 1
        row = table.add_row().cells
        _set_cell_width(row[0], widths[0]); _set_cell_vertical_center(row[0])
        _set_cell_width(row[1], widths[1]); _set_cell_vertical_center(row[1])
        _set_cell_width(row[2], widths[2]); _set_cell_vertical_center(row[2])

        # blank S.NO cell
        p0 = row[0].paragraphs[0]; p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _clear_para_spacing(p0)

        # topic name with bold "1.1" prefix + regular text
        p1 = row[1].paragraphs[0]; _clear_para_spacing(p1)
        _run(p1, f"{chap_no}.{topic_no}  ", 14, bold=True)
        _run(p1, str(topic), 14, bold=False)

        # page number, bold, centered, zero-padded
        page_str = str(page)
        if zero_pad:
            page_str = page_str.zfill(2)
        p2 = row[2].paragraphs[0]; p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _clear_para_spacing(p2)
        _run(p2, page_str, 18, bold=True)

    doc.save(output_path)
    return output_path


if __name__ == "__main__":
    # quick self-test with the biology data
    demo = [
        ("CELL THE UNIT OF LIFE", "Introduction to cell structure and cell theory", 1),
        ("CELL THE UNIT OF LIFE", "Comparing plant and animal cells", 21),
        ("CELL CYCLE: DIVISION", "Phases of the cell cycle and quiescent stage", 42),
        ("CELL CYCLE: DIVISION", "Mitosis: prophase and metaphase events", 58),
        ("CELL CYCLE: DIVISION", "Mitosis: anaphase, telophase, and chromosome dynamics", 76),
        ("CELL CYCLE: DIVISION", "Cytokinesis and Karyokynesis cell division in plants", 89),
        ("CELL CYCLE: DIVISION", "Meiosis i: prophase i substages and events", 100),
        ("CELL CYCLE: DIVISION", "Meiosis i: stages and chromosome behavior", 116),
        ("CELL CYCLE: DIVISION", "Meiosis ii: events and genetic outcomes", 131),
    ]
    out = generate_index(demo, "FINAL_INDEX.docx")
    print("wrote", out)
