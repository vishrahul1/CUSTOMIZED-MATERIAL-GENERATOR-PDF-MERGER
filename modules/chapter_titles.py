"""
chapter_titles.py
-----------------
Generate chapter divider pages — one full page per chapter — into a single
DOCX and convert it to a single PDF.

Each chapter page:
  • 21 × 27.8 cm page size  (11906 × 15761 DXA) — matches the project page size
  • twistedLines1 decorative border on all four edges
  • "CHAPTER – N" / "CHAPTER NAME"  centred, Calibri 36pt Bold Black

Pipeline role
-------------
These divider pages are generated AFTER the content has been merged and had its
headers/footers applied, then inserted at the start of each chapter (see
PDFMerger.insert_chapter_title_pages). They are intentionally NOT given the
class-name header or page-number footer, and they are unnumbered dividers, so
they do not disturb the content footer/index page numbering.

DOCX → PDF conversion uses a private, invisible Word instance (see
modules.word_convert), the same engine the index uses, because Word renders the
twistedLines1 art border correctly — without closing the user's open Word.
"""

import os

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from modules.word_convert import convert_docx_to_pdf


# ── Low-level helpers ───────────────────────────────────────────────────────

def _build_sectPr(page_break=False):
    """
    Return a fully-specified <w:sectPr> element.
    page_break=True  → used between chapters (nextPage section break)
    page_break=False → used as the final document section
    """
    sectPr = OxmlElement("w:sectPr")

    if page_break:
        t = OxmlElement("w:type")
        t.set(qn("w:val"), "nextPage")
        sectPr.append(t)

    # Page size: 21 × 27.8 cm
    pgSz = OxmlElement("w:pgSz")
    pgSz.set(qn("w:w"), "11906")
    pgSz.set(qn("w:h"), "15761")
    pgSz.set(qn("w:code"), "192")
    sectPr.append(pgSz)

    # Margins: 1 inch all sides
    pgMar = OxmlElement("w:pgMar")
    for attr, val in [("w:top", "1440"), ("w:right", "1440"),
                      ("w:bottom", "1440"), ("w:left", "1440"),
                      ("w:header", "708"), ("w:footer", "708"),
                      ("w:gutter", "0")]:
        pgMar.set(qn(attr), val)
    sectPr.append(pgMar)

    # Decorative border: twistedLines1 on all four sides
    pgBorders = OxmlElement("w:pgBorders")
    pgBorders.set(qn("w:offsetFrom"), "page")
    for side in ("w:top", "w:left", "w:bottom", "w:right"):
        b = OxmlElement(side)
        b.set(qn("w:val"),   "twistedLines1")
        b.set(qn("w:sz"),    "18")
        b.set(qn("w:space"), "24")
        b.set(qn("w:color"), "auto")
        pgBorders.append(b)
    sectPr.append(pgBorders)

    # Column grid
    cols = OxmlElement("w:cols")
    cols.set(qn("w:space"), "708")
    sectPr.append(cols)

    dg = OxmlElement("w:docGrid")
    dg.set(qn("w:linePitch"), "360")
    sectPr.append(dg)

    return sectPr


def _build_title_paragraph(doc, chapter_number, chapter_name, is_last=False):
    """
    Add one title paragraph (= one full chapter page) to *doc*.

    For every chapter except the last we embed a <w:sectPr> (nextPage break)
    inside the paragraph's <w:pPr> — the standard Word technique for per-page
    section breaks so each page keeps its own border. The last chapter uses the
    document-level sectPr (already configured).
    """
    chapter_number = str(chapter_number).strip()
    chapter_name   = str(chapter_name).strip().upper()

    para = doc.add_paragraph()
    pPr  = para._p.get_or_add_pPr()

    # Spacing — push the title roughly to the vertical middle of the page
    sp = OxmlElement("w:spacing")
    sp.set(qn("w:before"), "2320")
    sp.set(qn("w:after"),  "120")
    pPr.append(sp)

    # Centre alignment
    jc = OxmlElement("w:jc")
    jc.set(qn("w:val"), "center")
    pPr.append(jc)

    # Embed per-page sectPr (nextPage break) for all but the last chapter
    if not is_last:
        pPr.append(_build_sectPr(page_break=True))

    # Run ───────────────────────────────────────────────────────────────────
    run  = para.add_run()
    rPr  = run._r.get_or_add_rPr()

    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), "Calibri")
    rFonts.set(qn("w:hAnsi"), "Calibri")
    rFonts.set(qn("w:cs"),    "Calibri")
    rPr.insert(0, rFonts)

    rPr.append(OxmlElement("w:b"))
    rPr.append(OxmlElement("w:bCs"))

    color = OxmlElement("w:color")
    color.set(qn("w:val"), "000000")
    color.set(qn("w:themeColor"), "text1")
    rPr.append(color)

    for tag, val in [("w:sz", "72"), ("w:szCs", "72")]:
        e = OxmlElement(tag)
        e.set(qn("w:val"), val)
        rPr.append(e)

    lang = OxmlElement("w:lang")
    lang.set(qn("w:val"), "en-US")
    rPr.append(lang)

    # Text: line 1 + line break + line 2
    t1 = OxmlElement("w:t")
    t1.text = f"CHAPTER – {chapter_number}"
    run._r.append(t1)

    run._r.append(OxmlElement("w:br"))

    t2 = OxmlElement("w:t")
    t2.text = chapter_name
    run._r.append(t2)


# ── Public API ──────────────────────────────────────────────────────────────

def generate_chapter_titles_pdf(chapters, output_pdf_path):
    """
    Build one DOCX with a divider page per chapter and convert it to a single
    PDF whose pages are in the same order as *chapters*.

    Parameters
    ----------
    chapters : list of (number, name) tuples, in chapter order
               e.g. [(1, "Number System"), (2, "Real Numbers")]
    output_pdf_path : path to the .pdf to produce. The intermediate .docx is
                      written alongside it (same stem).

    Returns
    -------
    output_pdf_path
    """
    if not chapters:
        raise ValueError("[ChapterTitles] chapters list is empty")

    docx_path = os.path.splitext(output_pdf_path)[0] + ".docx"

    doc = Document()

    # Remove the default empty paragraph
    for p in list(doc.paragraphs):
        p._element.getparent().remove(p._element)

    # Configure the document-level sectPr (used by the last chapter's page)
    doc_sectPr = doc.sections[0]._sectPr
    for child in list(doc_sectPr):
        doc_sectPr.remove(child)
    for child in _build_sectPr(page_break=False):
        doc_sectPr.append(child)

    # One paragraph (page) per chapter
    for i, (num, name) in enumerate(chapters):
        is_last = (i == len(chapters) - 1)
        _build_title_paragraph(doc, num, name, is_last=is_last)

    doc.save(docx_path)
    print(f"[ChapterTitles] Generated chapter-titles DOCX at {docx_path}")

    # ── Convert .docx → .pdf via a private, invisible Word (win32com) ───────
    # Uses a dedicated Word instance so the user's open Word is never closed,
    # and avoids docx2pdf's tqdm (which crashes in no-console builds). Word
    # renders the twistedLines1 art border correctly.
    convert_docx_to_pdf(docx_path, output_pdf_path)

    print(f"[ChapterTitles] Generated chapter-titles PDF at {output_pdf_path}")
    return output_pdf_path
